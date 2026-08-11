"""Append the final real AD8232 clean-window analysis to the technical report."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TECHNICAL_REPORT_DOCX = ROOT / "Technical_Report" / "Real_Time_ECG_Monitor_Technical_Report.docx"
INPUT_DOCX = TECHNICAL_REPORT_DOCX
OUTPUT_DOCX = TECHNICAL_REPORT_DOCX
COMPARISON_JSON = ROOT / "results" / "real_ad8232_comparison" / "real_ad8232_session_comparison.json"
FIGURE_DIR = ROOT / "results" / "real_ad8232_comparison"
TECHNICAL_REPORT_FIGURE_DIR = ROOT / "docs" / "technical_report" / "figures"


def _fmt(value: float | None, digits: int = 2) -> str:
    if value is None:
        return "-"
    return f"{float(value):.{digits}f}"


def _pct(value: float | None) -> str:
    if value is None:
        return "-"
    return f"{100.0 * float(value):.1f}%"


def _rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi")
    if bidi is None:
        from docx.oxml import OxmlElement

        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


def _set_run_font(paragraph, size: int = 11, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold


def _add_paragraph(doc: Document, text: str, *, size: int = 11, bold: bool = False):
    paragraph = doc.add_paragraph(text)
    _rtl(paragraph)
    _set_run_font(paragraph, size=size, bold=bold)
    return paragraph


def _add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    _rtl(paragraph)
    _set_run_font(paragraph, size=16 if level == 1 else 13, bold=True)
    return paragraph


def _style_table(table) -> None:
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _rtl(paragraph)
                _set_run_font(paragraph, size=9)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    _style_table(table)


def _add_figure(
    doc: Document,
    image_name: str,
    caption: str,
    *,
    base_dir: Path = FIGURE_DIR,
    width_inches: float = 6.4,
) -> None:
    path = base_dir / image_name
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(caption_paragraph, size=9, bold=True)


def _remove_trailing_empty_paragraphs(doc: Document) -> None:
    while doc.paragraphs and not doc.paragraphs[-1].text.strip():
        element = doc.paragraphs[-1]._element
        element.getparent().remove(element)


def main() -> int:
    data = json.loads(COMPARISON_JSON.read_text(encoding="utf-8"))
    sessions = data["sessions"]
    segments = data["clean_segments"]

    doc = Document(INPUT_DOCX)
    _remove_trailing_empty_paragraphs(doc)
    doc.add_page_break()
    _add_heading(doc, "پیوست نهایی: تحلیل رکوردهای واقعی AD8232 بر اساس بازه‌های سالم", 1)
    _add_paragraph(
        doc,
        "این پیوست آموزشی و غیرتشخیصی است. به دلیل وجود artifact ناشی از جابه‌جایی دستگاه در بخشی از "
        "رکوردهای واقعی، تحلیل PQRST فقط روی سالم‌ترین پنجره‌های ۶ ثانیه‌ای انجام شد. معیار انتخاب شامل "
        "SQI، نبود lead-off، clipping پایین، پایداری RR، تعداد R-peak کافی و کامل بودن markerهای P/Q/R/S/T بود.",
    )

    _add_heading(doc, "setup سخت‌افزار نهایی", 2)
    _add_paragraph(
        doc,
        "در ثبت نهایی، Arduino فقط واحد acquisition بود و پردازش اصلی روی کامپیوتر انجام شد. اتصال نهایی "
        "با sketch فعلی چنین است: OUTPUT ماژول AD8232 به A5، پایه LO+ به D3، پایه LO- به D2، تغذیه به 3.3V "
        "و GND به زمین Arduino وصل شد. الکترودها مطابق برچسب‌های RA، LA و RL ماژول قرار گرفتند. این setup "
        "فقط برای آزمایش آموزشی استفاده شد و کاربرد پزشکی ندارد.",
    )
    _add_table(
        doc,
        ["پایه AD8232", "پایه Arduino", "نقش"],
        [
            ["OUTPUT", "A5", "ورودی آنالوگ ECG"],
            ["LO+", "D3", "تشخیص جدا شدن الکترود مثبت"],
            ["LO-", "D2", "تشخیص جدا شدن الکترود منفی"],
            ["VCC", "3.3V", "تغذیه ماژول"],
            ["GND", "GND", "زمین مشترک"],
        ],
    )
    _add_figure(
        doc,
        "hardware/electrode_placement_ra_la_rl.png",
        "شکل پیوست-۱: محل رایج الکترودهای RA/LA/RL در setup سه‌الکتروده آموزشی",
        base_dir=TECHNICAL_REPORT_FIGURE_DIR,
        width_inches=4.4,
    )
    _add_figure(
        doc,
        "hardware/ad8232_arduino_wiring_reference.png",
        "شکل پیوست-۲: اتصال Arduino Uno، ماژول AD8232 و سه الکترود",
        base_dir=TECHNICAL_REPORT_FIGURE_DIR,
        width_inches=5.5,
    )

    _add_heading(doc, "خلاصه acquisition", 2)
    _add_table(
        doc,
        ["فرد", "مدت", "نمونه معتبر", "نرخ نمونه‌برداری", "packet loss", "checksum", "SQI", "HR میانگین"],
        [
            [
                f"{s['sex_age']} ({s['subject_id']})",
                f"{_fmt(s['duration_s'])}s",
                str(s["valid_samples"]),
                f"{_fmt(s['estimated_sampling_rate_hz'], 3)}Hz",
                _pct(s["packet_loss_rate"]),
                str(s["checksum_errors"]),
                f"{s['overall_sqi_level']} ({_fmt(s['overall_sqi_score'])})",
                f"{_fmt(s['mean_hr_bpm'])} bpm",
            ]
            for s in sessions
        ],
    )

    _add_heading(doc, "بازه‌های سالم منتخب و خروجی PQRST", 2)
    _add_table(
        doc,
        ["فرد", "بازه", "زمان", "HR", "RR CV", "SQI", "خروجی rule-based"],
        [
            [
                f"{s['sex_age']} ({s['subject_id']})",
                s["condition"],
                f"{_fmt(s['start_s'], 1)}-{_fmt(s['end_s'], 1)}s",
                f"{_fmt(s['mean_hr_bpm'])} bpm",
                _fmt(s["rr_cv"], 3),
                f"{s['sqi_level']} ({_fmt(s['sqi_score'])})",
                s["rhythm_label"],
            ]
            for s in segments
        ],
    )

    doc.add_page_break()
    _add_heading(doc, "فاصله‌های زمانی تخمینی", 2)
    _add_table(
        doc,
        ["فرد", "بازه", "P-R peak", "QRS", "QT", "QTc", "P/Q/S/T visibility"],
        [
            [
                f"{s['subject_id']}",
                s["condition"],
                f"{_fmt(s['mean_p_to_r_ms'], 1)}ms",
                f"{_fmt(s['mean_qrs_ms'], 1)}ms",
                f"{_fmt(s['mean_qt_ms'], 1)}ms",
                f"{_fmt(s['mean_qtc_bazett_ms'], 1)}ms",
                f"P {_pct(s['p_visible_rate'])}, Q {_pct(s['q_visible_rate'])}, S {_pct(s['s_visible_rate'])}, T {_pct(s['t_visible_rate'])}",
            ]
            for s in segments
        ],
    )

    _add_paragraph(
        doc,
        "مقادیر P-R peak، QRS، QT و QTc در این پروژه تخمین‌های تک‌لید و آموزشی هستند و به‌عنوان تشخیص پزشکی "
        "یا اندازه‌گیری بالینی قطعی تفسیر نمی‌شوند. نتیجه اصلی این بخش آن است که سامانه توانسته بخش‌های خراب "
        "را کنار بگذارد و روی بخش‌های سالم markerگذاری و تحلیل عددی انجام دهد.",
    )

    _add_heading(doc, "نمودارهای نهایی", 2)
    _add_figure(doc, "real_subject_filtered_snippets.png", "شکل پیوست-۳: موج فیلترشده فقط در بازه‌های سالم منتخب")
    _add_figure(doc, "real_subject_gui_marker_snapshots.png", "شکل پیوست-۴: نمای شبیه GUI با markerهای P/Q/R/S/T روی بازه‌های سالم")
    _add_figure(doc, "real_subject_condition_hr_sqi.png", "شکل پیوست-۵: مقایسه HR و SQI در بازه‌های سالم")
    _add_figure(doc, "real_subject_acquisition_quality.png", "شکل پیوست-۶: شاخص‌های کیفیت acquisition در دو رکورد واقعی")

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
